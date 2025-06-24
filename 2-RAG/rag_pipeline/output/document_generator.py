import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generate DOCX documents from LLM responses"""
    
    def __init__(self):
        pass
        
    def create_docx(self, content: str, filename: str = "output.docx"):
        """Create a DOCX document from the generated content"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI Generated Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")  # Empty line
        
        # Process content and add to document
        paragraphs = content.split('\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Check if it's a heading
            if paragraph.startswith('#'):
                level = len(paragraph) - len(paragraph.lstrip('#'))
                heading_text = paragraph.lstrip('#').strip()
                doc.add_heading(heading_text, level)
            elif paragraph.startswith('•'):
                # Bullet point
                doc.add_paragraph(paragraph[1:].strip(), style='List Bullet')
            else:
                # Regular paragraph
                doc.add_paragraph(paragraph)
        
        # Save document
        doc.save(filename)
        logger.info(f"DOCX document saved as {filename}")
        
        return filename